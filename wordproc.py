import docx
import re

from commontools import CommonTools as ct
from dialog import Dialog as dlg
from wordsegmentation import WordSegmentation as ws

class WordProc():
    def __init__(self):
        """
        文档处理
        """
        self.word = None # 打开的文档对象
        self.file = None # 打开的文档路径
        self.content = {'序号': None, '通报名称': None, '发生日期': None, '问题描述': None, '责任人': None, '通报日期': None, '责任追究处理': None} # 要提取的文档内容
        
    def clearContent(self):
        """
        清空内容
        """
        self.word = None
        self.file = None
        self.content = {key: None for key in self.content.keys()}

    def openWord(self, file: str) -> bool:
        """
        打开文档

        Params:
            file (str): 文档路径

        Returns:
            成功与否状态
        """
        try:
            self.word = docx.Document(file)
        except Exception as e:
            self.clearContent()
            dlg.log(f'打开文件失败 {file}: {e}', dlg.ERROR)
            return False
        dlg.log(f'打开文件 {file}')
        self.file = file
        return True
        
    def getParagraphsCount(self) -> int:
        """
        获取文档段落数

        Returns:
            int: 段落数；-1 表示未打开文档
        """
        if self.word is not None:
            return len(self.word.paragraphs)
        else:
            return -1
        
    def getContentByIdx(self, paragraphIndex: int) -> tuple[bool, str]:
        """
        获取文档内容

        Params:
            paragraphIndex (int): 段落索引；-1 表示获取所有段落；其他值表示获取指定段落

        Returns:
            str: (是否成功，None/内容)
        """
        paragraphCount = self.getParagraphsCount()
        if paragraphCount == -1:
            dlg.log('未打开文档。请打开文档后再进行操作。', dlg.ERROR)
            return False, None
        elif paragraphIndex >= paragraphCount:
            dlg.log('段落索引超出范围。请检查段落索引后再进行操作。', dlg.ERROR)
            return False, None
        elif paragraphIndex == -1:
            return True, self.word.paragraphs
        else:
            return True, self.word.paragraphs[paragraphIndex].text.strip()
        
    def do(self) -> dict:
        """
        解析内容

        Returns:
            (True/False, dict): 成功与否, dict: {'序号': XX, '通报名称': XX, '发生日期': XX, '问题描述': XX, '责任人': XX, '通报日期': XX, '责任追究处理': xx}
        """
        # 序号
        status, serialNumberString = self.getContentByIdx(1)
        if not status:
            return False, None
        match = re.match(r'质量通报(.*)', serialNumberString)
        if match:
            self.content['序号'] = match.group(1)
        else:
            self.content['序号'] = None

        # 通报名称
        self.content['通报名称'] = ct.removeFileExtension((ct.getNameFromPath(self.file))).strip()

        # 发生日期 & 问题描述
        status, eventString = self.getContentByIdx(4)
        if not status:
            return False, None
        match = re.search(r'([^，。]+)，(.*)', eventString)
        if match:
            dateString, event = match.groups()
            self.content['发生日期'] = dateString
            self.content['问题描述'] = event

        # 责任人 & 责任追究处理
        isStart = False
        names = []
        responsibilityInvestigation = ''
        for idx in range(0, self.getParagraphsCount()):
            status, content = self.getContentByIdx(idx)
            if not status:
                return False, None
            if '责任追究处理' in content:
                isStart = True
                continue
            if '预防措施' in content:
                break
            if isStart:
                names += ws.do(content)
                responsibilityInvestigation += (content + '\n')
        names = list(dict.fromkeys(names)) # 去重
        self.content['责任人'] = names
        self.content['责任追究处理'] = responsibilityInvestigation

        # 通报日期
        status, reportDateString = self.getContentByIdx(self.getParagraphsCount() - 1)
        if not status:
            return False, None
        match = re.search(r'(\d{4})年(\d{1,2})月(\d{1,2})日', reportDateString)
        if match:
            year, month, day = match.groups()
            self.content['通报日期'] = f'{year}-{month}-{day}'

        return True, self.content